#!/usr/bin/env python3
import sys
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from io import BytesIO
import requests

def replace_text_in_paragraph(paragraph, placeholder, value):
    """Replace placeholder in paragraph, handling split runs and braces"""
    full_text = ''.join(run.text for run in paragraph.runs)

    # Try different formats: {{placeholder}}, {{Placeholder}}, {{PLACEHOLDER}}, and without braces
    patterns = [
        f'{{{{{placeholder}}}}}',  # {{Item.Name}}
        f'{{{{{placeholder.lower()}}}}}',  # {{item.name}}
        f'{{{{{placeholder.upper()}}}}}',  # {{ITEM.NAME}}
        placeholder,  # Item.Name (without braces)
        placeholder.lower(),  # item.name (without braces)
    ]

    found = False
    for pattern in patterns:
        if pattern in full_text:
            full_text = full_text.replace(pattern, str(value))
            found = True
            break

    if not found:
        return False

    # Clear all runs
    for run in paragraph.runs:
        run.text = ''

    # Add the new text to the first run
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)

    return True

def add_image_to_paragraph(paragraph, image_url):
    """Add image from URL to paragraph"""
    try:
        print(f"[DEBUG] Fetching image from URL...", file=sys.stderr)
        response = requests.get(image_url, timeout=10)
        print(f"[DEBUG] Response status: {response.status_code}", file=sys.stderr)
        if response.status_code == 200:
            print(f"[DEBUG] Image fetched successfully, size: {len(response.content)} bytes", file=sys.stderr)
            image_data = BytesIO(response.content)
            paragraph.add_run().add_picture(image_data, width=Inches(2))
            print(f"[DEBUG] Image inserted into document", file=sys.stderr)
            return True
        else:
            print(f"Warning: HTTP {response.status_code} - Could not fetch image from URL: {image_url[:100]}", file=sys.stderr)
    except Exception as e:
        print(f"Warning: Could not insert image {image_url[:100]}: {e}", file=sys.stderr)
    return False

def fill_template(template_path, output_path, data):
    """Fill a Word template with survey data"""

    print(f"Loading template from: {template_path}", file=sys.stderr)
    doc = Document(template_path)

    # Dictionary of replacements with both display name and actual placeholder
    replacements = {
        'Item.Name': data.get('areaName', 'N/A'),  # Room/Area name
        'Item.Client': data.get('clientName', 'N/A'),  # Vendor/Client name
        'Item.Area Name/ Room Number': data.get('areaName', 'N/A'),
        'Item.Area Size (Sqft)': data.get('areaSize', 'N/A'),
        'Item.Survey Date': data.get('surveyDate', 'N/A'),
        'Item.Survey Notes': data.get('surveyNotes', 'N/A'),
        'Item.Suggested System': data.get('recommendedSystem', 'N/A'),
        'Item.Notes Regarding Install': data.get('notes', 'N/A'),
    }

    print(f"Replacements to make: {replacements}", file=sys.stderr)
    replaced_count = 0

    # Replace placeholders in paragraphs
    print(f"Processing {len(doc.paragraphs)} paragraphs", file=sys.stderr)
    for idx, paragraph in enumerate(doc.paragraphs):
        para_text = ''.join(run.text for run in paragraph.runs)

        # Debug first 10 paragraphs with full details
        if idx < 10:
            print(f"Para {idx}: runs={len(paragraph.runs)}, text={repr(para_text[:150])}", file=sys.stderr)
            # Always show individual runs for paragraphs with runs
            if len(paragraph.runs) > 0:
                for run_idx, run in enumerate(paragraph.runs):
                    print(f"  Run {run_idx}: {repr(run.text)}", file=sys.stderr)

        for key, value in replacements.items():
            if replace_text_in_paragraph(paragraph, key, value):
                print(f"[OK] Replaced '{key}' = '{value}' in paragraph {idx}", file=sys.stderr)
                replaced_count += 1
            elif key in ['Item.Name', 'Item.Client'] and para_text.strip():
                # Debug logging for Item.Name and Item.Client
                if any(part in para_text for part in ['Item.Name', 'Item.Client', 'item.name', 'item.client']):
                    print(f"  ! Found '{key}' text but replacement failed in paragraph {idx}", file=sys.stderr)

    # Also replace in tables
    print(f"Processing {len(doc.tables)} tables", file=sys.stderr)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    cell_text = ''.join(run.text for run in paragraph.runs)
                    if cell_text.strip():
                        print(f"Table {table_idx}, row {row_idx}, cell {cell_idx}: {cell_text[:100]}", file=sys.stderr)
                    for key, value in replacements.items():
                        if replace_text_in_paragraph(paragraph, key, value):
                            print(f"Replaced '{key}' in table {table_idx}, row {row_idx}, cell {cell_idx}", file=sys.stderr)
                            replaced_count += 1

    print(f"Total text replacements made: {replaced_count}", file=sys.stderr)

    # Skipping header/footer processing to focus on main content
    # print(f"Checking {len(doc.sections)} sections for headers/footers", file=sys.stderr)
    print(f"Total replacements: {replaced_count}", file=sys.stderr)

    # Comprehensive search for ALL text in the document
    print(f"Comprehensive search for remaining placeholders", file=sys.stderr)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def search_and_replace_in_element(element, replacements):
        """Recursively search and replace in all text elements"""
        count = 0
        element_count = 0
        t_elems = list(element.iter(qn('w:t')))

        print(f"Total text elements found: {len(t_elems)}", file=sys.stderr)

        # Show first 20 text elements found
        for i, t_elem in enumerate(t_elems[:20]):
            if t_elem.text:
                print(f"Text element {i+1}: {repr(t_elem.text[:80])}", file=sys.stderr)

        # Handle split placeholders by merging adjacent elements
        i = 0
        while i < len(t_elems):
            t_elem = t_elems[i]
            if t_elem.text and '{{' in t_elem.text:
                # Potential start of a placeholder, collect the next elements
                collected = [t_elem.text]
                j = i + 1
                last_closing = j - 1
                while j < len(t_elems) and j < i + 10:  # Look ahead up to 10 elements
                    next_text = t_elems[j].text or ''
                    collected.append(next_text)
                    if '}}' in next_text:
                        # Found the closing braces
                        merged = ''.join(collected)

                        # Skip image and scan placeholders - let them be handled by dedicated image insertion
                        if 'Item.Images' in merged or 'Item.Scans' in merged or 'Images of Area' in merged or 'Scans of Area' in merged:
                            i = j
                            break

                        for key, value in replacements.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in merged:
                                # Replace the entire merged text, keeping structure
                                new_text = merged.replace(placeholder, str(value))
                                print(f"[DEBUG] Replacing split placeholder '{key}': '{merged}' -> '{new_text}'", file=sys.stderr)
                                print(f"[DEBUG] Elements involved: {i} to {j} (count: {j-i+1})", file=sys.stderr)
                                t_elem.text = new_text
                                # Only clear the elements that are purely part of the placeholder (not spacing)
                                # For elements 2-4 ('{{',' Item.Name', '}}'), we need to be careful:
                                # Keep the first element with replacement, don't touch others to preserve spacing
                                for k in range(i + 1, j + 1):
                                    old_text = t_elems[k].text or ''
                                    # Only clear if it's clearly placeholder text, not spacing
                                    if old_text.strip() in ['Item.Name', 'Item.Client', '}}', '{{']:
                                        print(f"[DEBUG] Clearing element {k}: '{old_text}'", file=sys.stderr)
                                        t_elems[k].text = ''
                                print(f"[OK] Replaced '{key}' (split across {j-i+1} elements)", file=sys.stderr)
                                count += 1
                                i = j
                                break
                        break
                    j += 1
            i += 1

        return count

    found_count = search_and_replace_in_element(doc.element, replacements)
    if found_count > 0:
        print(f"Found and replaced {found_count} placeholders in comprehensive search", file=sys.stderr)
        replaced_count += found_count

    # Handle images - look for Item.Images of Area placeholder
    images_added = 0
    if data.get('images') and len(data['images']) > 0:
        print(f"Attempting to add {len(data['images'])} images", file=sys.stderr)
        print(f"Image URLs: {data['images']}", file=sys.stderr)
        # Find the paragraph with Item.Images of Area placeholder
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = ''.join(run.text for run in paragraph.runs)
            if 'Item.Images' in para_text or 'Images of Area' in para_text:
                print(f"Found Images placeholder at paragraph {para_idx}: {para_text[:100]}", file=sys.stderr)
                # Clear the placeholder text
                for run in paragraph.runs:
                    run.text = ''
                # Add all images to this paragraph
                for idx, image_url in enumerate(data['images'][:5]):  # Limit to 5 images
                    print(f"Attempting to add image {idx + 1}: {image_url}", file=sys.stderr)
                    if add_image_to_paragraph(paragraph, image_url):
                        images_added += 1
                        print(f"Successfully added image {idx + 1}/{len(data['images'])}", file=sys.stderr)
                break

        # Also check in tables
        if images_added == 0:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            para_text = ''.join(run.text for run in paragraph.runs)
                            if 'Item.Images' in para_text or 'Images of Area' in para_text:
                                print(f"Found Images placeholder in table {table_idx}, row {row_idx}, cell {cell_idx}: {para_text[:100]}", file=sys.stderr)
                                # Clear the placeholder text
                                for run in paragraph.runs:
                                    run.text = ''
                                # Add all images to this paragraph
                                for idx, image_url in enumerate(data['images'][:5]):  # Limit to 5 images
                                    print(f"Attempting to add image {idx + 1} to table: {image_url}", file=sys.stderr)
                                    if add_image_to_paragraph(paragraph, image_url):
                                        images_added += 1
                                        print(f"Successfully added image {idx + 1}/{len(data['images'])} to table", file=sys.stderr)
                                break

    print(f"Total images added: {images_added}", file=sys.stderr)

    # Handle scans - look for Item.Scans of Area placeholder
    scans_added = 0
    if data.get('scans') and len(data['scans']) > 0:
        print(f"Attempting to add {len(data['scans'])} scans", file=sys.stderr)
        # Find the paragraph with Item.Scans of Area placeholder
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = ''.join(run.text for run in paragraph.runs)
            if 'Item.Scans' in para_text or 'Scans of Area' in para_text:
                print(f"Found Scans placeholder at paragraph {para_idx}: {para_text[:100]}", file=sys.stderr)
                # Clear the placeholder text
                for run in paragraph.runs:
                    run.text = ''
                # Add all scans to this paragraph
                for idx, scan_url in enumerate(data['scans'][:5]):  # Limit to 5 scans
                    if add_image_to_paragraph(paragraph, scan_url):
                        scans_added += 1
                        print(f"Added scan {idx + 1}/{len(data['scans'])}", file=sys.stderr)
                break

        # Also check in tables
        if scans_added == 0:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            para_text = ''.join(run.text for run in paragraph.runs)
                            if 'Item.Scans' in para_text or 'Scans of Area' in para_text:
                                print(f"Found Scans placeholder in table {table_idx}, row {row_idx}, cell {cell_idx}", file=sys.stderr)
                                # Clear the placeholder text
                                for run in paragraph.runs:
                                    run.text = ''
                                # Add all scans to this paragraph
                                for idx, scan_url in enumerate(data['scans'][:5]):  # Limit to 5 scans
                                    if add_image_to_paragraph(paragraph, scan_url):
                                        scans_added += 1
                                break

    print(f"Total scans added: {scans_added}", file=sys.stderr)

    # Save the document
    print(f"Saving document to: {output_path}", file=sys.stderr)
    doc.save(output_path)
    print(f"Template fill complete", file=sys.stderr)
    return True

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(json.dumps({'error': 'Usage: fill_template.py <template_path> <output_path> <json_data>'}))
        sys.exit(1)

    template_path = sys.argv[1]
    output_path = sys.argv[2]
    log_file = 'fill_template.log'

    # Redirect stderr to also write to log file
    class DualWriter:
        def __init__(self, file1, file2):
            self.file1 = file1
            self.file2 = file2
        def write(self, msg):
            self.file1.write(msg)
            self.file2.write(msg)
            self.file2.flush()
        def flush(self):
            self.file1.flush()
            self.file2.flush()

    with open(log_file, 'a') as lf:
        lf.write(f"\n=== Starting fill_template.py ===\n")
        original_stderr = sys.stderr
        sys.stderr = DualWriter(original_stderr, lf)

        try:
            # Parse JSON data from argument
            if len(sys.argv) > 3:
                data = json.loads(sys.argv[3])
            else:
                data = json.load(sys.stdin)

            print(f"Starting template fill", file=sys.stderr)
            fill_template(template_path, output_path, data)
            result = json.dumps({'success': True, 'output': output_path})
            print(result)
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            print(error_msg, file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)
            print(json.dumps({'error': str(e)}), file=sys.stderr)
            sys.exit(1)
        finally:
            sys.stderr = original_stderr
